from datetime import date, datetime
from io import BytesIO
from collections import Counter
from django.http import HttpResponse
from django.db.models import Count, Q
from openpyxl import Workbook
from openpyxl.styles import Font, PatternFill, Border, Side, Alignment
from openpyxl.chart import LineChart, BarChart, Reference
from openpyxl.utils import get_column_letter
from openpyxl.formatting.rule import ColorScaleRule, CellIsRule
from rest_framework import viewsets
from rest_framework.decorators import api_view
from rest_framework.response import Response
from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Pt
from .models import *
from .serializers import *

SEMESTER_START = date(2026, 2, 10)  # пары начались с нижней недели

class DepartmentViewSet(viewsets.ModelViewSet):
    queryset = Department.objects.all().order_by('name')
    serializer_class = DepartmentSerializer

class StudyGroupViewSet(viewsets.ModelViewSet):
    serializer_class = StudyGroupSerializer
    def get_queryset(self):
        qs = StudyGroup.objects.select_related('department').all().order_by('number')
        dep = self.request.query_params.get('department')
        if dep: qs = qs.filter(department_id=dep)
        return qs

class PersonViewSet(viewsets.ModelViewSet):
    serializer_class = PersonSerializer
    def get_queryset(self):
        qs = Person.objects.select_related('group','department').all().order_by('full_name')
        role = self.request.query_params.get('role')
        group = self.request.query_params.get('group')
        department = self.request.query_params.get('department')
        subgroup = self.request.query_params.get('subgroup')
        if role: qs = qs.filter(role=role)
        if group: qs = qs.filter(group_id=group)
        if department: qs = qs.filter(Q(department_id=department)|Q(group__department_id=department))
        if subgroup: qs = qs.filter(subgroup=subgroup)
        return qs
    def perform_create(self, serializer):
        obj = serializer.save()
        if obj.role == 'student' and obj.is_monitor and obj.group_id:
            Person.objects.filter(role='student', group_id=obj.group_id, is_monitor=True).exclude(id=obj.id).update(is_monitor=False)
    def perform_update(self, serializer):
        obj = serializer.save()
        if obj.role == 'student' and obj.is_monitor and obj.group_id:
            Person.objects.filter(role='student', group_id=obj.group_id, is_monitor=True).exclude(id=obj.id).update(is_monitor=False)

class SubjectViewSet(viewsets.ModelViewSet):
    serializer_class = SubjectSerializer
    def get_queryset(self):
        qs = Subject.objects.select_related('department').all().order_by('title')
        dep = self.request.query_params.get('department')
        if dep: qs = qs.filter(department_id=dep)
        return qs

class LessonViewSet(viewsets.ModelViewSet):
    serializer_class = LessonSerializer
    def get_queryset(self):
        qs = Lesson.objects.select_related('subject','teacher','group','group__department').all().order_by('day_of_week','time')
        teacher = self.request.query_params.get('teacher')
        group = self.request.query_params.get('group')
        subject = self.request.query_params.get('subject')
        department = self.request.query_params.get('department')
        subgroup = self.request.query_params.get('subgroup')
        if teacher: qs = qs.filter(teacher_id=teacher)
        if group: qs = qs.filter(group_id=group)
        if subject: qs = qs.filter(subject_id=subject)
        if department: qs = qs.filter(group__department_id=department)
        qs = apply_subgroup_to_lessons(qs, subgroup)
        return qs

class AttendanceViewSet(viewsets.ModelViewSet):
    serializer_class = AttendanceSerializer
    def get_queryset(self):
        qs = Attendance.objects.select_related('student','lesson','lesson__subject','lesson__teacher','lesson__group','lesson__group__department').all().order_by('-date','lesson__time')
        return filter_attendance(qs, self.request.query_params)

def week_type_for(d: date):
    weeks = (d - SEMESTER_START).days // 7
    return 'n' if weeks % 2 == 0 else 'v'

def lesson_subgroup_number_text(value):
    value = value or ''
    if 'гр.1' in value or value == '1': return '1'
    if 'гр.2' in value or value == '2': return '2'
    return ''

def apply_subgroup_to_lessons(qs, subgroup):
    if subgroup:
        return qs.filter(Q(subgroup='') | Q(subgroup__icontains=f'гр.{subgroup}'))
    return qs

def filter_attendance(qs, params):
    department = params.get('department')
    group = params.get('group')
    subject = params.get('subject')
    teacher = params.get('teacher')
    date_from = params.get('date_from')
    date_to = params.get('date_to')
    subgroup = params.get('subgroup')
    if department: qs = qs.filter(lesson__group__department_id=department)
    if group: qs = qs.filter(lesson__group_id=group)
    if subject: qs = qs.filter(lesson__subject_id=subject)
    if teacher: qs = qs.filter(lesson__teacher_id=teacher)
    if date_from: qs = qs.filter(date__gte=date_from)
    if date_to: qs = qs.filter(date__lte=date_to)
    if subgroup: qs = qs.filter(student__subgroup=subgroup)
    return qs

@api_view(['POST'])
def login(request):
    login = request.data.get('login','').strip()
    password = request.data.get('password','').strip()
    user = Person.objects.filter(login=login, password=password).first()
    if not user:
        return Response({'error':'Неверный логин или пароль'}, status=401)
    return Response(PersonSerializer(user).data)

@api_view(['GET'])
def today_lessons(request):
    raw = request.GET.get('date')
    d = datetime.strptime(raw, '%Y-%m-%d').date() if raw else date.today()
    wt = week_type_for(d)
    lessons = Lesson.objects.filter(day_of_week=d.weekday()).filter(Q(week_type=wt)|Q(week_type='both')).select_related('subject','teacher','group')
    role = request.GET.get('role')
    user_id = request.GET.get('user_id')
    group = request.GET.get('group')
    subject = request.GET.get('subject')
    if group: lessons = lessons.filter(group_id=group)
    if subject: lessons = lessons.filter(subject_id=subject)
    if role == 'student' and user_id:
        user = Person.objects.filter(id=user_id).first()
        if user and user.group:
            lessons = lessons.filter(group=user.group)
            lessons = apply_subgroup_to_lessons(lessons, user.subgroup)
    if role == 'teacher' and user_id:
        lessons = lessons.filter(teacher_id=user_id)
    return Response({'date': d, 'week_type': wt, 'week_label': 'нижняя' if wt=='n' else 'верхняя', 'lessons': LessonSerializer(lessons, many=True).data})

@api_view(['POST'])
def mark_attendance(request):
    lesson_id = request.data.get('lesson')
    marked_by_id = request.data.get('marked_by')
    attend_date = request.data.get('date') or str(date.today())
    statuses = request.data.get('statuses', [])
    saved = []
    lesson = Lesson.objects.filter(id=lesson_id).first()
    lesson_sg = lesson_subgroup_number_text(lesson.subgroup if lesson else '')
    for row in statuses:
        student = Person.objects.filter(id=row['student']).first()
        if lesson_sg and student and student.subgroup != lesson_sg:
            continue
        obj, _ = Attendance.objects.update_or_create(
            student_id=row['student'], lesson_id=lesson_id, date=attend_date,
            defaults={'status': row['status'], 'marked_by_id': marked_by_id, 'reason': row.get('reason','')}
        )
        saved.append(obj)
    return Response({'saved': len(saved), 'items': AttendanceSerializer(saved, many=True).data})

@api_view(['GET'])
def stats(request):
    qs = filter_attendance(Attendance.objects.all(), request.GET)
    total = qs.count(); present = qs.filter(status='present').count(); late = qs.filter(status='late').count(); absent = qs.filter(status='absent').count()
    return Response({'total_marks': total, 'present': present, 'late': late, 'absent': absent, 'attendance_percent': round(((present+late)/total*100) if total else 0, 1)})

@api_view(['GET'])
def student_attendance(request, pk):
    qs = Attendance.objects.filter(student_id=pk).select_related('lesson','lesson__subject','lesson__teacher').order_by('-date','lesson__time')
    return Response(AttendanceSerializer(qs, many=True).data)

@api_view(['GET'])
def semester_report(request):
    qs = filter_attendance(Attendance.objects.select_related('student','lesson','lesson__subject','lesson__group'), request.GET)
    return Response(AttendanceSerializer(qs.order_by('date','lesson__time'), many=True).data)

@api_view(['GET'])
def export_report(request):
    qs = filter_attendance(
        Attendance.objects.select_related(
            'student', 'student__group', 'lesson', 'lesson__subject', 'lesson__teacher',
            'lesson__group', 'lesson__group__department'
        ),
        request.GET
    )

    group_id = request.GET.get('group')
    department_id = request.GET.get('department')
    subject_id = request.GET.get('subject')
    teacher_id = request.GET.get('teacher')
    subgroup = request.GET.get('subgroup')
    date_from = request.GET.get('date_from')
    date_to = request.GET.get('date_to')

    students = Person.objects.select_related('group', 'group__department').filter(role='student')
    if group_id:
        students = students.filter(group_id=group_id)
    if department_id:
        students = students.filter(Q(department_id=department_id) | Q(group__department_id=department_id))
    if subgroup:
        students = students.filter(subgroup=subgroup)

    # В отчёт попадают только студенты, по которым есть отметки после применения фильтров.
    marked_student_ids = set(qs.values_list('student_id', flat=True).distinct())
    if marked_student_ids:
        students = students.filter(id__in=marked_student_ids)
    students = students.order_by('group__number', 'full_name')

    total = qs.count()
    present = qs.filter(status='present').count()
    absent = qs.filter(status='absent').count()
    attendance_percent = round((present / total * 100), 1) if total else 0

    group = StudyGroup.objects.filter(id=group_id).first() if group_id else None
    subject = Subject.objects.filter(id=subject_id).first() if subject_id else None
    teacher = Person.objects.filter(id=teacher_id).first() if teacher_id else None

    doc = Document()
    section = doc.sections[0]
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width, section.page_height = section.page_height, section.page_width
    section.top_margin = Cm(1.4)
    section.bottom_margin = Cm(1.4)
    section.left_margin = Cm(1.2)
    section.right_margin = Cm(1.2)

    title = 'Сводный отчёт по посещаемости студентов'
    if group:
        title += f' группы {group.number}'
    heading = doc.add_heading(title, level=1)
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER

    generated_at = datetime.now().strftime('%d.%m.%Y %H:%M')
    doc.add_paragraph(f'Дата формирования: {generated_at}')

    filters = []
    if department_id:
        dep = Department.objects.filter(id=department_id).first()
        if dep: filters.append(f'кафедра/направление: {dep.name}')
    if group: filters.append(f'группа: {group.number}')
    if subgroup: filters.append(f'подгруппа: {subgroup}')
    if subject: filters.append(f'дисциплина: {subject.title}')
    if teacher: filters.append(f'преподаватель: {teacher.full_name}')
    if date_from or date_to: filters.append(f'период: {date_from or "начало"} — {date_to or "текущая дата"}')
    doc.add_paragraph('Применённые фильтры: ' + ('; '.join(filters) if filters else 'не заданы, отчёт сформирован по всем доступным данным.'))

    kpi = doc.add_table(rows=2, cols=4)
    kpi.style = 'Table Grid'
    for i, h in enumerate(['Всего отметок', 'Присутствий', 'Пропусков', 'Средняя посещаемость']):
        kpi.rows[0].cells[i].text = h
    for i, v in enumerate([str(total), str(present), str(absent), f'{attendance_percent}%']):
        kpi.rows[1].cells[i].text = v

    doc.add_paragraph('')
    doc.add_heading('Детализация по студентам', level=2)
    table = doc.add_table(rows=1, cols=11)
    table.style = 'Table Grid'
    headers = [
        '№', 'ФИО студента', 'Группа', 'Подгр.', 'Всего занятий', 'Посещено',
        'Пропусков', '% посещ.', 'Полных дней пропусков', 'Дисциплины с пропусками', 'Причины / примечания'
    ]
    for i, header in enumerate(headers):
        table.rows[0].cells[i].text = header

    for i, student in enumerate(students, 1):
        student_qs = qs.filter(student=student)
        student_total = student_qs.count()
        student_present = student_qs.filter(status='present').count()
        student_absent_qs = student_qs.filter(status='absent')
        student_absent = student_absent_qs.count()
        student_percent = round((student_present / student_total * 100), 1) if student_total else 0

        full_days = student_absent_qs.values('date').annotate(c=Count('id')).filter(c__gte=2).count()
        subject_counter = Counter(student_absent_qs.values_list('lesson__subject__title', flat=True))
        subjects_text = '; '.join([f'{name} — {count}' for name, count in subject_counter.items()]) or 'нет'
        reasons = sorted(set([r for r in student_absent_qs.values_list('reason', flat=True) if r]))
        reasons_text = '; '.join(reasons) if reasons else 'не указаны'

        row = table.add_row().cells
        row[0].text = str(i)
        row[1].text = student.full_name
        row[2].text = student.group.number if student.group else ''
        row[3].text = student.subgroup or ''
        row[4].text = str(student_total)
        row[5].text = str(student_present)
        row[6].text = str(student_absent)
        row[7].text = f'{student_percent}%'
        row[8].text = str(full_days)
        row[9].text = subjects_text
        row[10].text = reasons_text

    doc.add_paragraph('')
    doc.add_heading('Вывод по отчёту', level=2)
    doc.add_paragraph(
        f'По выбранной выборке зафиксировано {total} отметок, из них {present} присутствий и {absent} пропусков. '
        f'Средний показатель посещаемости составляет {attendance_percent}%. '
        'Студенты с наибольшим числом пропусков требуют дополнительного контроля со стороны преподавателя, куратора или деканата.'
    )
    doc.add_paragraph('Ответственное лицо __________________ / __________________')

    for paragraph in doc.paragraphs:
        for run in paragraph.runs:
            run.font.name = 'Times New Roman'
            run.font.size = Pt(12)

    bio = BytesIO()
    doc.save(bio)
    bio.seek(0)
    resp = HttpResponse(bio.read(), content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
    resp['Content-Disposition'] = 'attachment; filename="attendance_report_filtered.docx"'
    return resp

@api_view(['GET'])
def analytics(request):
    qs = filter_attendance(Attendance.objects.select_related('student','student__group','lesson','lesson__subject','lesson__teacher','lesson__group','lesson__group__department'), request.GET)
    total = qs.count(); present = qs.filter(status='present').count(); late = qs.filter(status='late').count(); absent = qs.filter(status='absent').count()
    effective_present = present + late

    by_subject=[]
    for subject_id in qs.values_list('lesson__subject_id', flat=True).distinct():
        subject = Subject.objects.filter(id=subject_id).first()
        if not subject: continue
        sub_qs = qs.filter(lesson__subject=subject); t=sub_qs.count(); p=sub_qs.filter(status__in=['present','late']).count(); a=sub_qs.filter(status='absent').count(); l=sub_qs.filter(status='late').count()
        by_subject.append({'subject':subject.title,'total':t,'present':p,'absent':a,'late':l,'percent':round(p/t*100,1) if t else 0})
    by_subject.sort(key=lambda x: x['percent'])

    by_student=[]
    for st in Person.objects.filter(role='student', id__in=qs.values_list('student_id', flat=True).distinct()).select_related('group').order_by('full_name'):
        sq=qs.filter(student=st); t=sq.count(); p=sq.filter(status__in=['present','late']).count(); a=sq.filter(status='absent').count(); l=sq.filter(status='late').count(); percent=round(p/t*100,1) if t else 0
        status = 'Группа риска' if percent < 50 else 'Предупреждение' if percent < 70 else 'Норма'
        by_student.append({'student':st.full_name,'group': st.group.number if st.group else '', 'subgroup':st.subgroup,'total':t,'present':p,'absent':a,'late':l,'percent':percent,'risk':status})
    by_student.sort(key=lambda x: (x['percent'], -x['absent']))

    by_group=[]
    for g in StudyGroup.objects.filter(id__in=qs.values_list('lesson__group_id', flat=True).distinct()).select_related('department'):
        gq=qs.filter(lesson__group=g); t=gq.count(); p=gq.filter(status__in=['present','late']).count(); a=gq.filter(status='absent').count(); l=gq.filter(status='late').count()
        students = [r for r in by_student if r['group'] == g.number]
        best = max(students, key=lambda x: x['percent'])['student'] if students else '—'
        worst = min(students, key=lambda x: x['percent'])['student'] if students else '—'
        by_group.append({'group':g.number,'department': g.department.name if g.department else '', 'total':t,'present':p,'absent':a,'late':l,'percent':round(p/t*100,1) if t else 0, 'best_student':best, 'worst_student':worst})
    by_group.sort(key=lambda x: x['percent'], reverse=True)

    by_department=[]
    for dep in Department.objects.filter(groups__lessons__attendance__in=qs).distinct():
        dq=qs.filter(lesson__group__department=dep); t=dq.count(); p=dq.filter(status__in=['present','late']).count(); a=dq.filter(status='absent').count()
        by_department.append({'department':dep.name,'total':t,'absent':a,'percent':round(p/t*100,1) if t else 0})

    by_teacher=[]
    for teacher in Person.objects.filter(role='teacher', lessons__attendance__in=qs).distinct().order_by('full_name'):
        tq=qs.filter(lesson__teacher=teacher)
        lessons_total=tq.values('lesson_id','date').distinct().count()
        filled=tq.values('lesson_id','date').distinct().count()
        by_teacher.append({'teacher':teacher.full_name,'lessons':lessons_total,'filled_journals':filled})

    by_week=[]
    weekly = {}
    for a in qs:
        key = f'{a.date.isocalendar().year}-W{a.date.isocalendar().week:02d}'
        weekly.setdefault(key, {'total':0,'present':0})
        weekly[key]['total'] += 1
        if a.status in ['present','late']:
            weekly[key]['present'] += 1
    for key in sorted(weekly):
        t=weekly[key]['total']; p=weekly[key]['present']
        by_week.append({'week':key,'percent':round(p/t*100,1) if t else 0})

    heatmap=[]
    daily={}
    for a in qs:
        daily.setdefault(str(a.date), {'total':0,'present':0})
        daily[str(a.date)]['total'] += 1
        if a.status in ['present','late']:
            daily[str(a.date)]['present'] += 1
    for d in sorted(daily):
        t=daily[d]['total']; p=daily[d]['present']; percent=round(p/t*100,1) if t else 0
        heatmap.append({'date':d,'percent':percent,'level':'high' if percent >= 85 else 'mid' if percent >= 70 else 'low'})

    absences=[]
    for a in qs.filter(status='absent').order_by('-date','student__full_name')[:500]:
        absences.append({
            'date': a.date.strftime('%d.%m.%Y'),
            'student': a.student.full_name,
            'group': a.student.group.number if a.student.group else (a.lesson.group.number if a.lesson.group else ''),
            'subject': a.lesson.subject.title if a.lesson and a.lesson.subject else '',
            'reason': a.reason or 'Не указана',
        })

    return Response({'total':total,'present':present,'late':late,'absent':absent,'percent':round(effective_present/total*100,1) if total else 0,'by_subject':by_subject,'by_student':by_student,'by_group':by_group,'by_department':by_department,'by_teacher':by_teacher,'by_week':by_week,'heatmap':heatmap,'absences':absences})

@api_view(['GET'])
def semester_matrix(request):
    teacher_id=request.GET.get('teacher')
    group_id=request.GET.get('group')
    subject_id=request.GET.get('subject')
    subgroup=request.GET.get('subgroup')
    lessons=Lesson.objects.select_related('subject','group','teacher').all()
    if teacher_id: lessons=lessons.filter(teacher_id=teacher_id)
    if group_id: lessons=lessons.filter(group_id=group_id)
    if subject_id: lessons=lessons.filter(subject_id=subject_id)
    lessons=apply_subgroup_to_lessons(lessons, subgroup)
    params = request.GET.copy()
    att=filter_attendance(Attendance.objects.filter(lesson__in=lessons).select_related('student','lesson','lesson__subject').order_by('date','lesson__time'), params)
    cols=[]; seen=set()
    for a in att:
        key=f'{a.date}|{a.lesson_id}'
        if key not in seen:
            seen.add(key); cols.append({'key':key,'date':a.date,'time':a.lesson.time,'subject':a.lesson.subject.title})
    rows=[]
    students=Person.objects.filter(role='student', attendance__lesson__in=lessons).distinct().order_by('full_name')
    if group_id: students=students.filter(group_id=group_id)
    if subgroup: students=students.filter(subgroup=subgroup)
    for st in students:
        mp={}
        for a in att.filter(student=st):
            mp[f'{a.date}|{a.lesson_id}']='+' if a.status=='present' else '-'
        rows.append({'student':st.full_name,'subgroup':st.subgroup,'cells':[mp.get(c['key'],'') for c in cols], 'absent':sum(1 for v in mp.values() if v=='-')})
    return Response({'columns':cols,'rows':rows})

@api_view(['GET'])
def student_subject_summary(request, pk):
    out=[]
    for s in Subject.objects.all():
        qs=Attendance.objects.filter(student_id=pk, lesson__subject=s).select_related('lesson','lesson__teacher').order_by('date','lesson__time')
        t=qs.count(); p=qs.filter(status='present').count(); a=qs.filter(status='absent').count()
        if t:
            absences=[]
            for item in qs.filter(status='absent'):
                absences.append({
                    'date': item.date,
                    'time': item.lesson.time,
                    'lesson_type': item.lesson.lesson_type,
                    'teacher': item.lesson.teacher.full_name if item.lesson.teacher else '',
                    'room': item.lesson.room,
                })
            out.append({'subject':s.title,'total':t,'present':p,'absent':a,'percent':round(p/t*100,1),'absences':absences})
    return Response(out)

def _sheet_title(ws, title, subtitle=''):
    ws.merge_cells(start_row=1, start_column=1, end_row=1, end_column=8)
    ws.cell(1, 1, title).font = Font(size=16, bold=True, color='FFFFFF')
    ws.cell(1, 1).fill = PatternFill('solid', fgColor='0057B8')
    ws.cell(1, 1).alignment = Alignment(horizontal='center')
    if subtitle:
        ws.merge_cells(start_row=2, start_column=1, end_row=2, end_column=8)
        ws.cell(2, 1, subtitle).font = Font(italic=True, color='4B5563')

def _write_table(ws, start_row, headers, rows, percent_cols=None):
    percent_cols = percent_cols or []
    header_fill = PatternFill('solid', fgColor='DCEBFF')
    thin = Side(style='thin', color='B8C7DD')
    for col, header in enumerate(headers, 1):
        c = ws.cell(start_row, col, header)
        c.font = Font(bold=True, color='0F2A50')
        c.fill = header_fill
        c.border = Border(top=thin, bottom=thin, left=thin, right=thin)
        c.alignment = Alignment(horizontal='center', vertical='center', wrap_text=True)
    for r, row in enumerate(rows, start_row + 1):
        for col, value in enumerate(row, 1):
            c = ws.cell(r, col, value)
            c.border = Border(top=thin, bottom=thin, left=thin, right=thin)
            c.alignment = Alignment(vertical='top', wrap_text=True)
            if col in percent_cols:
                c.number_format = '0.0%'
    for col_idx in range(1, ws.max_column + 1):
        values = [ws.cell(row_idx, col_idx).value for row_idx in range(1, ws.max_row + 1)]
        max_len = max((len(str(value)) if value is not None else 0) for value in values)
        ws.column_dimensions[get_column_letter(col_idx)].width = min(max(max_len + 2, 12), 34)

@api_view(['GET'])
def export_report_xlsx(request):
    qs = filter_attendance(
        Attendance.objects.select_related('student','student__group','lesson','lesson__subject','lesson__teacher','lesson__group','lesson__group__department'),
        request.GET
    ).order_by('date','lesson__time','student__full_name')

    wb = Workbook()
    wb.remove(wb.active)
    generated = datetime.now().strftime('%d.%m.%Y %H:%M')

    total = qs.count(); present = qs.filter(status='present').count(); late = qs.filter(status='late').count(); absent = qs.filter(status='absent').count()
    percent = round(((present + late) / total), 4) if total else 0

    ws = wb.create_sheet('Сводка')
    _sheet_title(ws, 'Сводный отчет по посещаемости', f'Сформировано: {generated}')
    _write_table(ws, 4, ['Показатель', 'Значение'], [
        ['Всего отметок', total], ['Посещено', present], ['Опоздания', late], ['Пропущено', absent], ['Средняя посещаемость', percent]
    ], percent_cols=[2])
    ws['B8'].number_format = '0.0%'

    # 1. Отчет по посещаемости группы / рейтинг студентов / опоздания / риск.
    student_rows = []
    for st in Person.objects.filter(role='student', id__in=qs.values_list('student_id', flat=True).distinct()).select_related('group').order_by('group__number','full_name'):
        sq = qs.filter(student=st); t=sq.count(); p=sq.filter(status__in=['present','late']).count(); a=sq.filter(status='absent').count(); l=sq.filter(status='late').count(); pct_val=round(p/t,4) if t else 0
        status = 'Группа риска' if pct_val < .5 else 'Предупреждение' if pct_val < .7 else 'Норма'
        student_rows.append([st.full_name, st.group.number if st.group else '', st.subgroup, p, a, l, pct_val, status])
    ws = wb.create_sheet('Студенты')
    _sheet_title(ws, 'Отчет по посещаемости студентов', 'Группа, рейтинг, опоздания и статус риска')
    _write_table(ws, 4, ['Студент','Группа','Подгр.','Посещено','Пропущено','Опоздания','% посещаемости','Статус'], student_rows, percent_cols=[7])
    if len(student_rows) > 0:
        ws.conditional_formatting.add(f'G5:G{4+len(student_rows)}', ColorScaleRule(start_type='num', start_value=0.4, start_color='F8696B', mid_type='num', mid_value=0.7, mid_color='FFEB84', end_type='num', end_value=1, end_color='63BE7B'))

    ranking = sorted(student_rows, key=lambda r: r[6], reverse=True)
    ws = wb.create_sheet('Рейтинг')
    _sheet_title(ws, 'Рейтинг посещаемости студентов')
    _write_table(ws, 4, ['Место','Студент','Группа','% посещаемости'], [[i+1, r[0], r[1], r[6]] for i, r in enumerate(ranking)], percent_cols=[4])

    risk = [r for r in student_rows if r[7] != 'Норма']
    ws = wb.create_sheet('Группа риска')
    _sheet_title(ws, 'Студенты группы риска', 'Ниже 70% — предупреждение, ниже 50% — группа риска')
    _write_table(ws, 4, ['Студент','Группа','% посещаемости','Статус'], [[r[0], r[1], r[6], r[7]] for r in risk], percent_cols=[3])

    # 2. Отчет по дисциплинам.
    subj_rows = []
    for subject in Subject.objects.filter(id__in=qs.values_list('lesson__subject_id', flat=True).distinct()).order_by('title'):
        sq=qs.filter(lesson__subject=subject); lessons=sq.values('lesson_id','date').distinct().count(); t=sq.count(); p=sq.filter(status__in=['present','late']).count()
        subj_rows.append([subject.title, lessons, round(p/t,4) if t else 0, sq.filter(status='absent').count(), sq.filter(status='late').count()])
    ws = wb.create_sheet('Дисциплины')
    _sheet_title(ws, 'Отчет по дисциплинам')
    _write_table(ws, 4, ['Дисциплина','Всего занятий','Средняя посещаемость','Пропуски','Опоздания'], subj_rows, percent_cols=[3])

    # 3. Журнал пропусков.
    absence_rows = []
    for a in qs.filter(status='absent'):
        absence_rows.append([a.date.strftime('%d.%m.%Y'), a.student.full_name, a.student.group.number if a.student.group else '', a.lesson.subject.title, a.reason or 'Не указана'])
    ws = wb.create_sheet('Журнал пропусков')
    _sheet_title(ws, 'Журнал пропусков студентов')
    _write_table(ws, 4, ['Дата','Студент','Группа','Дисциплина','Причина'], absence_rows)

    # 5. Факультет/кафедра.
    dep_rows = []
    for dep in Department.objects.filter(groups__lessons__attendance__in=qs).distinct().order_by('name'):
        dq=qs.filter(lesson__group__department=dep); t=dq.count(); p=dq.filter(status__in=['present','late']).count()
        dep_rows.append([dep.name, round(p/t,4) if t else 0, dq.filter(status='absent').count()])
    ws = wb.create_sheet('Факультеты')
    _sheet_title(ws, 'Отчет по факультету / направлению')
    _write_table(ws, 4, ['Факультет / кафедра','Средняя посещаемость','Количество пропусков'], dep_rows, percent_cols=[2])

    # 7. Преподаватели.
    teacher_rows = []
    for teacher in Person.objects.filter(role='teacher', lessons__attendance__in=qs).distinct().order_by('full_name'):
        tq=qs.filter(lesson__teacher=teacher)
        teacher_rows.append([teacher.full_name, tq.values('lesson_id','date').distinct().count(), tq.values('lesson_id','date').distinct().count()])
    ws = wb.create_sheet('Преподаватели')
    _sheet_title(ws, 'Отчет по преподавателям')
    _write_table(ws, 4, ['Преподаватель','Проведено занятий','Заполнено журналов'], teacher_rows)

    # 8/12. Сравнение групп + семестр.
    group_rows = []
    for g in StudyGroup.objects.filter(id__in=qs.values_list('lesson__group_id', flat=True).distinct()).order_by('number'):
        gq=qs.filter(lesson__group=g); t=gq.count(); p=gq.filter(status__in=['present','late']).count(); students=[r for r in student_rows if r[1]==g.number]
        best=max(students, key=lambda x:x[6])[0] if students else '—'; worst=min(students, key=lambda x:x[6])[0] if students else '—'
        group_rows.append([g.number, round(p/t,4) if t else 0, best, worst])
    ws = wb.create_sheet('Сравнение групп')
    _sheet_title(ws, 'Сводный отчет за семестр и сравнение групп')
    _write_table(ws, 4, ['Группа','Средняя посещаемость','Лучший студент','Худшая посещаемость'], group_rows, percent_cols=[2])
    if len(group_rows) > 0:
        chart = BarChart(); chart.title = 'ТОП групп по посещаемости'; chart.y_axis.title = '%'; chart.x_axis.title = 'Группа'
        data = Reference(ws, min_col=2, min_row=4, max_row=4+len(group_rows)); cats = Reference(ws, min_col=1, min_row=5, max_row=4+len(group_rows))
        chart.add_data(data, titles_from_data=True); chart.set_categories(cats); ws.add_chart(chart, 'F4')

    # 9. Динамика по неделям.
    weekly = {}
    for a in qs:
        key=f'{a.date.isocalendar().year}-W{a.date.isocalendar().week:02d}'
        weekly.setdefault(key, {'total':0,'present':0})
        weekly[key]['total'] += 1
        if a.status in ['present','late']: weekly[key]['present'] += 1
    week_rows = [[k, round(v['present']/v['total'],4) if v['total'] else 0] for k, v in sorted(weekly.items())]
    ws = wb.create_sheet('Динамика')
    _sheet_title(ws, 'График посещаемости по неделям')
    _write_table(ws, 4, ['Неделя','% посещаемости'], week_rows, percent_cols=[2])
    if len(week_rows) > 1:
        chart = LineChart(); chart.title='Динамика посещаемости'; chart.y_axis.title='%'; chart.x_axis.title='Неделя'
        data=Reference(ws,min_col=2,min_row=4,max_row=4+len(week_rows)); cats=Reference(ws,min_col=1,min_row=5,max_row=4+len(week_rows))
        chart.add_data(data,titles_from_data=True); chart.set_categories(cats); ws.add_chart(chart,'E4')

    # 10. Тепловая карта.
    daily = {}
    for a in qs:
        key=a.date.strftime('%d.%m.%Y')
        daily.setdefault(key, {'total':0,'present':0})
        daily[key]['total'] += 1
        if a.status in ['present','late']: daily[key]['present'] += 1
    heat_rows = [[k, round(v['present']/v['total'],4) if v['total'] else 0] for k, v in sorted(daily.items())]
    ws = wb.create_sheet('Тепловая карта')
    _sheet_title(ws, 'Тепловая карта посещаемости по датам')
    _write_table(ws, 4, ['Дата','% посещаемости'], heat_rows, percent_cols=[2])
    if heat_rows:
        ws.conditional_formatting.add(f'B5:B{4+len(heat_rows)}', ColorScaleRule(start_type='num', start_value=0.4, start_color='F8696B', mid_type='num', mid_value=0.75, mid_color='FFEB84', end_type='num', end_value=1, end_color='63BE7B'))

    bio = BytesIO(); wb.save(bio); bio.seek(0)
    resp = HttpResponse(bio.read(), content_type='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet')
    resp['Content-Disposition'] = 'attachment; filename="attendance_analytics_report.xlsx"'
    return resp
